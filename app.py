import sys
import os
os.environ["STREAMLIT_WATCH_MODULES"] = "false"
if 'STREAMLIT_RUN_MUTED' not in os.environ:
    os.environ['STREAMLIT_RUN_MUTED'] = 'true'
# Lấy đường dẫn tuyệt đối của thư mục gốc
ROOT_DIR = os.path.dirname(os.path.abspath(__file__))

# Thêm các thư mục cần thiết vào PYTHONPATH
sys.path.insert(0, ROOT_DIR)
sys.path.insert(0, os.path.join(ROOT_DIR, "src"))
sys.path.insert(0, os.path.join(ROOT_DIR, "dictionary"))

# Fixes for PyTorch and Streamlit compatibility issues
os.environ["STREAMLIT_WATCH_MODULES"] = "false"

# Disable Streamlit's file watcher to avoid PyTorch class issues
if 'STREAMLIT_RUN_MUTED' not in os.environ:
    os.environ['STREAMLIT_RUN_MUTED'] = 'true'

import streamlit as st
import json
from word_analysis import analyze_word_sentiment
from sentence_analysis import SentenceAnalyzer
from dictionary.dict_manager import DictionaryManager
import pandas as pd
from io import StringIO
from docx import Document
import datetime

def add_word(word, score):
    # Validate input
    if not (-1 <= score <= 1):
        st.error("Điểm số phải trong khoảng [-1, 1]")
        return False
    
    # Cập nhật từ điển sử dụng DictionaryManager
    success = dict_manager.add_word_to_sentiment_dict(word, score)
    
    if success:
        # Trigger re-render
        st.session_state.update_trigger = not st.session_state.get('update_trigger', False)
        # Hiển thị thông báo
        st.success(f"Đã thêm từ '{word}' với điểm {score}")
        return True
    else:
        st.error(f"Có lỗi khi lưu từ '{word}' vào từ điển")
        return False


def reload_dictionaries():
    """Reload all dictionaries from JSON files"""
    global dict_manager
    dict_manager.load_all()

# Khởi tạo Dictionary Manager
dict_manager = DictionaryManager(os.path.join(ROOT_DIR, "dictionary"))

# Lấy các từ điển
SENTIMENT_DICT = dict_manager.get_sentiment_dict()
NEGATION_WORDS = dict_manager.get_negation_words()
INTENSIFIER_WORDS = dict_manager.get_intensifier_words()
DIMINISHER_WORDS = dict_manager.get_diminisher_words()
COMPOUND_WORDS = dict_manager.get_compound_words()
PROVERBS = dict_manager.get_proverbs()
PUNCTUATION_ANALYSIS = dict_manager.get_punctuation_analysis()

# Cấu hình trang
st.set_page_config(
    page_title="Phân tích cảm xúc tiếng Việt",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Khởi tạo session state
if 'sentence_analyzer' not in st.session_state:
    st.session_state.sentence_analyzer = SentenceAnalyzer()

# Lưu lịch sử phân tích vào session state
if 'analysis_history' not in st.session_state:
    st.session_state.analysis_history = []

# Hàm lưu lịch sử phân tích
def get_status_from_sentiment(sentence_result):
    # Xác định trạng thái dựa trên cảm xúc chính
    if isinstance(sentence_result, dict):
        main = sentence_result.get('main_sentiment')
        if main:
            if 'tích cực' in main.lower() or 'positive' in main.lower():
                return 'Tích cực'
            elif 'tiêu cực' in main.lower() or 'negative' in main.lower():
                return 'Tiêu cực'
            elif 'trung lập' in main.lower() or 'neutral' in main.lower():
                return 'Trung lập'
    return 'Không xác định'

def save_analysis_history(text, word_result, sentence_result):
    status = get_status_from_sentiment(sentence_result)
    st.session_state.analysis_history.append({
        'timestamp': datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'text': text,
        'word_result': word_result,
        'sentence_result': sentence_result,
        'status': status
    })

# Hàm xuất báo cáo ra file docx
def export_report_docx(history):
    doc = Document()
    doc.add_heading('Báo cáo lịch sử phân tích cảm xúc', 0)
    for entry in history:
        doc.add_heading(f"Thời gian: {entry['timestamp']}", level=1)
        doc.add_paragraph(f"Văn bản: {entry['text']}")
        doc.add_paragraph(f"Trạng thái: {entry.get('status', 'Không xác định')}")
        doc.add_heading('Kết quả phân tích từ:', level=2)
        for word in entry['word_result']['words']:
            doc.add_paragraph(f"- {word['word']}: {word['score']} ({word.get('type','')})")
        doc.add_heading('Kết quả phân tích câu:', level=2)
        doc.add_paragraph(str(entry['sentence_result']))
    buffer = StringIO()
    doc.save('report.docx')
    return 'report.docx'

# Hàm giả lập huấn luyện lại mô hình
@st.cache_data(show_spinner=False)
def retrain_model(train_data):
    # Placeholder: Thực tế sẽ gọi script huấn luyện, lưu model mới
    import time
    time.sleep(2)
    return 'Huấn luyện lại mô hình thành công!'

# Tạo tabs cho ứng dụng
tab1, tab2, tab3 = st.tabs(["Phân tích cảm xúc", "Huấn luyện mô hình", "Thống kê & Báo cáo"])

with tab1:
    # Tiêu đề và mô tả
    st.title("Phân tích cảm xúc tiếng Việt")
    st.markdown("""
        Ứng dụng phân tích cảm xúc văn bản tiếng Việt, hỗ trợ:
        - Phân tích từng từ và cụm từ
        - Xác định cảm xúc tích cực, tiêu cực, trung lập
        - Hiển thị kết quả trực quan
    """)

    # Ô nhập văn bản
    text = st.text_area("Nhập văn bản cần phân tích:", height=100, help="Nhập văn bản tiếng Việt cần phân tích cảm xúc")

    # Xử lý phân tích
    if st.button("Phân tích cảm xúc"):
        if not text.strip():
            st.warning("⚠️ Vui lòng nhập văn bản cần phân tích!")
        else:
            try:
                with st.spinner('Đang phân tích...'):
                    # Phân tích từ ngữ
                    word_analysis = analyze_word_sentiment(text)
                    
                    
                    # Phân tích câu văn
                    sentence_analysis = st.session_state.sentence_analyzer.analyze_sentence(text, word_analysis)
                    
                    # Hiển thị kết quả
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.subheader("📝 Phân tích từng từ")
                        
                        # Tạo bảng phân tích từ
                        word_data = []
                        total_score = 0
                        
                        # Lưu điểm số trước khi xử lý dấu câu
                        pre_punct_score = 0
                        punct_effect = 0
                        
                        for word in word_analysis["words"]:
                            # Xác định loại từ
                            word_type_display = {
                                "positive": "Tích cực 🟢",
                                "negative": "Tiêu cực 🔴",
                                "neutral": "Trung lập ⚪",
                                "negation": "Phủ định ❌",
                                "intensifier": "Tăng cường ⬆️",
                                "diminisher": "Giảm nhẹ ⬇️",
                                "compound": "Cụm từ 🔤",
                                "Tích cực 🟢": "Tích cực 🟢",
                                "Tiêu cực 🔴": "Tiêu cực 🔴",
                                "Trung lập ⚪": "Trung lập ⚪"
                            }.get(word["type"], word["type"])
                            
                            # Xác định loại cảm xúc dựa vào điểm số
                            sentiment_type = "Trung lập ⚪"
                            if word["score"] > 0:
                                sentiment_type = "Tích cực 🟢"
                            elif word["score"] < 0:
                                sentiment_type = "Tiêu cực 🔴"
                                
                            # Lưu các thông tin về loại từ vào phần mô tả
                            type_info = word_type_display
                            if "description" not in word or not word["description"]:
                                word["description"] = type_info
                            
                            # Thêm mô tả nếu có
                            description = word.get("description", "")
                            
                            # Nếu từ là biểu tượng cảm xúc, lấy mô tả từ PUNCTUATION_ANALYSIS
                            if word["word"] in PUNCTUATION_ANALYSIS:
                                emoji_info = PUNCTUATION_ANALYSIS[word["word"]]
                                if "description" in emoji_info:
                                    description = emoji_info["description"]
                                
                                # Hiển thị loại cảm xúc dựa trên effect
                                if "effect" in emoji_info:
                                    if emoji_info["effect"] == "positive":
                                        sentiment_type = "Tích cực 🟢"
                                    elif emoji_info["effect"] == "negative":
                                        sentiment_type = "Tiêu cực 🔴"
                                
                                # Hiển thị giá trị chính xác từ PUNCTUATION_ANALYSIS nếu có
                                if "value" in emoji_info:
                                    word["score"] = emoji_info["value"]  # Cập nhật điểm mới
                            
                            # Cộng điểm cho tất cả các loại từ và dấu câu
                            total_score += word["score"]
                            
                            # Thêm vào bảng dữ liệu
                            word_data.append({
                                "Từ/Cụm từ": word["word"],
                                "Loại": sentiment_type,
                                "Điểm": f"{word['score']:.2f}",
                                "Mô tả": description
                            })
                        
                        # Cập nhật tổng điểm trong word_analysis
                        word_analysis['score'] = total_score
                        
                        st.table(word_data)
                        
                        # Hiển thị tổng điểm
                        st.metric("Tổng điểm", f"{word_analysis['score']:.2f}")
                        
                        # Hiển thị thông tin về ảnh hưởng của dấu câu nếu có
                        if "pre_punct_score" in word_analysis and "punct_effect" in word_analysis:
                            pre_punct_score = word_analysis["pre_punct_score"]
                            punct_effect = word_analysis["punct_effect"]
                            
                            col1a, col1b = st.columns(2)
                            with col1a:
                                st.metric("Điểm trước dấu câu", f"{pre_punct_score:.2f}")
                            with col1b:
                                st.metric("Ảnh hưởng dấu câu", f"{punct_effect:.2f}", 
                                         delta_color="normal" if punct_effect == 0 else "off")
                    
                    with col2:
                        st.subheader("📊 Phân tích toàn câu")
                        
                        # Tính tổng điểm tuyệt đối của tất cả các từ có điểm khác 0
                        total_abs = sum(abs(word["score"]) for word in word_analysis["words"] if word["score"] != 0)
                        
                        # Chỉ tính các từ có điểm khác 0
                        meaningful_words = [word for word in word_analysis["words"] if word["score"] != 0]
                        
                        if total_abs > 0:
                            # Tính tổng điểm tích cực và tiêu cực
                            positive_score = sum(word["score"] for word in word_analysis["words"] if word["score"] > 0)
                            negative_score = abs(sum(word["score"] for word in word_analysis["words"] if word["score"] < 0))
                            
                            sentiment_data = {
                                "Tích cực 🟢": positive_score,
                                "Tiêu cực 🔴": negative_score
                            }
                            
                            # Chỉ thêm trung lập nếu có từ trung lập có ý nghĩa
                            neutral_score = sum(1 for word in meaningful_words 
                                             if abs(word["score"]) <= 0.1) / len(meaningful_words) if meaningful_words else 0
                            
                            if neutral_score > 0:
                                sentiment_data["Trung lập ⚪"] = neutral_score
                            
                            # Chuẩn hóa để tổng = 1
                            total = sum(sentiment_data.values())
                            if total > 0:
                                for key in sentiment_data:
                                    sentiment_data[key] = sentiment_data[key] / total
                        else:
                            sentiment_data = {
                                "Tích cực 🟢": 0,
                                "Tiêu cực 🔴": 0,
                                "Trung lập ⚪": 1
                            }
                        
                        # Xác định cảm xúc chính và tính độ tin cậy
                        if total_abs > 0:
                            # Tính tỷ lệ điểm số
                            positive_ratio = sum(word["score"] for word in word_analysis["words"] if word["score"] > 0) / total_abs
                            negative_ratio = abs(sum(word["score"] for word in word_analysis["words"] if word["score"] < 0)) / total_abs
                            
                            # Xác định cảm xúc chính dựa trên tổng điểm thực tế
                            # Sử dụng total_score từ word_analysis từ result
                            final_score = word_analysis["score"]
                            if final_score < -0.1:  # Ngưỡng tiêu cực
                                main_sentiment = "Tiêu cực 🔴"
                                # Độ tin cậy dựa trên mức độ tiêu cực và tỷ lệ từ tiêu cực
                                confidence = min(abs(final_score) * 50, 100) * negative_ratio
                            elif final_score > 0.1:  # Ngưỡng tích cực
                                main_sentiment = "Tích cực 🟢"
                                # Độ tin cậy dựa trên mức độ tích cực và tỷ lệ từ tích cực
                                confidence = min(final_score * 50, 100) * positive_ratio
                            else:
                                main_sentiment = "Trung lập ⚪"
                                # Độ tin cậy trung lập dựa trên mức độ cân bằng giữa tích cực và tiêu cực
                                balance_factor = 1 - abs(positive_ratio - negative_ratio)
                                confidence = balance_factor * 80  # Giảm độ tin cậy tối đa xuống 80% cho trung lập
                        else:
                            # Nếu không có điểm số có ý nghĩa, thì là trung lập với độ tin cậy vừa phải
                            main_sentiment = "Trung lập ⚪"
                            confidence = 70.0  # Giảm độ tin cậy xuống 70% khi không có từ có ý nghĩa
                        
                        # Giới hạn độ tin cậy trong khoảng 0-100%
                        confidence = min(max(confidence, 0.0), 100.0)
                        
                        # Hiển thị kết quả chính
                        st.metric(
                            "Cảm xúc chính",
                            main_sentiment,
                            f"Độ tin cậy: {confidence:.1f}%"
                        )
                        
                        # Vẽ biểu đồ
                        chart = st.bar_chart(sentiment_data)
                        
                        # Thống kê chi tiết
                        st.markdown("**Chi tiết phân tích:**")
                        stats = {
                            "Số từ phân tích": len([w for w in word_analysis["words"] 
                                                  if w["type"] not in ["intensifier", "diminisher", "negation"]]),
                            "Từ tích cực": sum(1 for w in word_analysis["words"] if w["score"] > 0),
                            "Từ tiêu cực": sum(1 for w in word_analysis["words"] if w["score"] < 0),
                            "Từ trung lập": sum(1 for w in word_analysis["words"] 
                                              if w["score"] == 0 and w["type"] not in ["intensifier", "diminisher", "negation"])
                        }
                        st.write(stats)
                        
                save_analysis_history(text, word_analysis, sentence_analysis)
            except Exception as e:
                st.error(f"❌ Có lỗi xảy ra: {str(e)}")
                st.error("Vui lòng thử lại hoặc liên hệ hỗ trợ!")

with tab2:
    st.title("Huấn luyện mô hình")
    st.markdown("""
        Thêm từ ngữ mới hoặc biểu tượng cảm xúc vào từ điển để cải thiện khả năng phân tích cảm xúc.
    """)
    
    # Tạo các tab cho các loại huấn luyện khác nhau
    train_tab1, train_tab2, train_tab3, train_tab4 = st.tabs(["⚡ Thêm từ đơn mới", "😊 Thêm biểu tượng cảm xúc", "👥 Thêm cụm từ", "📜 Thêm thành ngữ/tục ngữ"])

with train_tab1:
    st.subheader("Thêm từ ngữ mới vào từ điển")
    # Form nhập từ mới
    with st.form("add_word_form"):
        new_word = st.text_input("Nhập từ mới (CHỈ NHẬP MỘT TỪ đơn lẻ):", placeholder="Ví dụ: tuyệt")
        sentiment_score = st.slider("Điểm số:", min_value=-1.0, max_value=1.0, value=0.0, step=0.1, key="word_score_slider")
        # Tự động xác định loại cảm xúc dựa trên điểm số
        from config import SENTIMENT_CONFIG

        def determine_sentiment_type(score):
            # Đảm bảo phân loại chính xác với ngưỡng 0.1
            if score >= 0.1:
                return 'Tích cực'
            elif score <= -0.1:
                return 'Tiêu cực'
            else:
                return 'Trung lập'
        
        # Sử dụng giá trị hiện tại của slider
        sentiment_type = determine_sentiment_type(sentiment_score)
        st.write(f'**Loại tự động:** {sentiment_type}')
        submit_word = st.form_submit_button("Thêm từ")
    if submit_word and new_word.strip():
        # Kiểm tra xem input có phải từ đơn không
        if len(new_word.split()) > 1:
            st.error(f"❌ '{new_word}' chứa nhiều từ! Vui lòng chỉ nhập MỘT TỪ đơn lẻ duy nhất. Nếu muốn thêm cụm từ, hãy sử dụng tab 'Thêm cụm từ mới'.")
        else:
            try:
                # Kiểm tra xem từ đã tồn tại chưa
                if new_word in SENTIMENT_DICT:
                    st.warning(f"⚠️ Từ '{new_word}' đã tồn tại trong từ điển!")
                else:
                    # Thêm từ mới vào từ điển bằng dictionary manager
                    if dict_manager.add_word_to_sentiment_dict(new_word, sentiment_score):
                        # Cập nhật biến toàn cục
                        SENTIMENT_DICT = dict_manager.get_sentiment_dict()
                        
                        st.success(f"✅ Đã thêm từ '{new_word}' với điểm số {sentiment_score} (loại: {sentiment_type}) vào từ điển!")
                        st.info("ℹ️ Từ mới đã được thêm vào và sẵn sàng sử dụng ngay lập tức.")
                    else:
                        st.error("❌ Không thể thêm từ vào từ điển!")
            except Exception as e:
                st.error(f"❌ Có lỗi xảy ra khi cập nhật từ điển: {str(e)}")

# Hiển thị danh sách từ hiện có
    st.subheader("Danh sách từ hiện có")
    search_word = st.text_input("Tìm kiếm từ:", placeholder="Nhập từ cần tìm...")
    try:
        reload_dictionaries()
        updated_dict = dict_manager.get_sentiment_dict()
    except Exception:
        updated_dict = SENTIMENT_DICT
    if updated_dict:
        filtered_dict = {k: v for k, v in updated_dict.items() if not search_word or search_word.lower() in k.lower()}
        st.info(f"Tìm thấy {len(filtered_dict)} từ" + (f" chứa '{search_word}'" if search_word else ""))
        if filtered_dict:
            word_table = []
            for word, score in filtered_dict.items():
                sentiment = "Tích cực 🟢" if score >= 0.1 else "Tiêu cực 🔴" if score <= -0.1 else "Trung lập ⚪"
                word_table.append({"Từ/Cụm từ": word, "Điểm số": f"{score:.1f}", "Loại": sentiment})
            st.table(word_table)
    else:
        st.error("Không thể tải danh sách từ từ SENTIMENT_DICT.")

with train_tab2:
    st.subheader("Thêm biểu tượng cảm xúc mới")
    # Form nhập biểu tượng cảm xúc mới
    with st.form("add_emoji_form"):
        new_emoji = st.text_input("Nhập biểu tượng cảm xúc:", placeholder="Ví dụ: 😊")
        emoji_value = st.slider("Giá trị cảm xúc:", min_value=-1.0, max_value=1.0, value=0.0, step=0.1)
        # Tự động xác định loại cảm xúc dựa trên điểm số
        if emoji_value >= 0.1:
            auto_effect = "positive"
            effect_label = "Tích cực"
        elif emoji_value <= -0.1:
            auto_effect = "negative"
            effect_label = "Tiêu cực"
        else:
            auto_effect = "neutral"
            effect_label = "Trung lập"
        st.write(f'**Loại tự động:** {effect_label}')
        description = st.text_input("Mô tả:", value=f"Emoji {effect_label.lower()} thể hiện cảm xúc {effect_label.lower()}")
        submit_emoji = st.form_submit_button("Thêm biểu tượng cảm xúc")
    if submit_emoji and new_emoji.strip():
        try:
            # Kiểm tra xem biểu tượng đã tồn tại chưa
            if new_emoji in PUNCTUATION_ANALYSIS:
                st.warning(f"⚠️ Biểu tượng '{new_emoji}' đã tồn tại trong từ điển!")
            else:
                # Thêm biểu tượng mới vào từ điển bằng dictionary manager, dùng auto_effect
                if dict_manager.add_emoji_to_punctuation(new_emoji, auto_effect, emoji_value, description):
                    # Cập nhật biến toàn cục
                    PUNCTUATION_ANALYSIS = dict_manager.get_punctuation_analysis()
                    
                    st.success(f"✅ Đã thêm biểu tượng '{new_emoji}' vào từ điển!")
                    # Thông báo thành công
                    st.info("ℹ️ Biểu tượng đã được thêm vào và sẵn sàng sử dụng ngay lập tức.")
                else:
                    st.error("❌ Không thể thêm biểu tượng vào từ điển!")
        except Exception as e:
            st.error(f"❌ Có lỗi xảy ra khi cập nhật từ điển: {str(e)}")
    # Hiển thị danh sách biểu tượng cảm xúc hiện có
    st.subheader("Danh sách biểu tượng cảm xúc hiện có")
    try:
        reload_dictionaries()
        updated_punctuation = dict_manager.get_punctuation_analysis()
        # Hiển thị từ điển dưới dạng bảng
        emoji_table = []
        for emoji, info in updated_punctuation.items():
            # Xử lý cả hai trường hợp: 'value' và 'multiplier'
            if "value" in info:
                score = info["value"]
            elif "multiplier" in info:
                score = info["multiplier"]
            else:
                score = 0.0
                
            sentiment = "Tích cực 🟢" if info["effect"] == "positive" else "Tiêu cực 🔴" if info["effect"] == "negative" else "Trung lập ⚪"
            emoji_table.append({"Biểu tượng": emoji, "Điểm số": f"{score:.1f}", "Loại": sentiment})
        st.table(emoji_table)
    except Exception as e:
        st.error(f"❌ Có lỗi xảy ra khi cập nhật từ điển: {str(e)}")

with train_tab3:
    st.subheader("Thêm cụm từ mới")
    # Form nhập cụm từ mới
    with st.form("add_compound_form"):
        new_compound = st.text_input("Nhập cụm từ mới:", placeholder="Ví dụ: không ngừng nhớ")
        compound_score = st.slider("Điểm số:", min_value=-1.0, max_value=1.0, value=0.0, step=0.1, key="compound_score_slider")
        
        # Tự động xác định loại cảm xúc dựa trên điểm số
        compound_type = "Tích cực" if compound_score >= 0.1 else "Tiêu cực" if compound_score <= -0.1 else "Trung lập"
        st.write(f'**Loại tự động:** {compound_type}')
        
        submit_compound = st.form_submit_button("Thêm cụm từ")
    if submit_compound and new_compound.strip():
        try:
            # Kiểm tra xem cụm từ đã tồn tại chưa
            if new_compound in COMPOUND_WORDS:
                st.warning(f"⚠️ Cụm từ '{new_compound}' đã tồn tại trong từ điển!")
            else:
                # Thêm cụm từ mới vào từ điển bằng dictionary manager
                if dict_manager.add_compound_word(new_compound, compound_score):
                    # Cập nhật biến toàn cục
                    COMPOUND_WORDS = dict_manager.get_compound_words()
                    
                    st.success(f"✅ Đã thêm cụm từ '{new_compound}' với điểm số {compound_score} (loại: {compound_type}) vào từ điển!")
                    st.info("ℹ️ Cụm từ mới đã được thêm vào và sẵn sàng sử dụng ngay lập tức.")
                else:
                    st.error("❌ Không thể thêm cụm từ vào từ điển!")
        except Exception as e:
            st.error(f"❌ Có lỗi xảy ra khi cập nhật từ điển: {str(e)}")
    search_compound = st.text_input("Tìm kiếm cụm từ:", placeholder="Nhập cụm từ cần tìm...", key="search_compound")
    try:
        reload_dictionaries()
        updated_compounds = dict_manager.get_compound_words()
    except Exception:
        updated_compounds = COMPOUND_WORDS
    # Lọc và hiển thị từ điển
    filtered_compounds = {k: v for k, v in updated_compounds.items() if not search_compound or search_compound.lower() in k.lower()}
    # Hiển thị số lượng cụm từ
    st.info(f"Tìm thấy {len(filtered_compounds)} cụm từ" + (f" chứa '{search_compound}'" if search_compound else ""))
    # Hiển thị từ điển dưới dạng bảng
    if filtered_compounds:
        compound_table = []
        for word, score in filtered_compounds.items():
            sentiment = "Tích cực 🟢" if score >= 0.1 else "Tiêu cực 🔴" if score <= -0.1 else "Trung lập ⚪"
            compound_table.append({"Cụm từ": word, "Điểm số": f"{score:.1f}", "Loại": sentiment})
        st.table(compound_table)

# Tab thêm thành ngữ/tục ngữ
with train_tab4:
    st.subheader("Thêm thành ngữ/tục ngữ mới")
    # Form nhập thành ngữ/tục ngữ mới
    with st.form("add_proverb_form"):
        new_proverb = st.text_input("Nhập thành ngữ/tục ngữ mới:", placeholder="Ví dụ: ăn quả nhớ kẻ trồng cây")
        proverb_score = st.slider("Điểm số:", min_value=-1.0, max_value=1.0, value=0.0, step=0.1, key="proverb_score_slider")
        
        # Tự động xác định loại cảm xúc dựa trên điểm số
        proverb_type = determine_sentiment_type(proverb_score)
        st.write(f'**Loại tự động:** {proverb_type}')
        
        submit_proverb = st.form_submit_button("Thêm thành ngữ/tục ngữ")
    if submit_proverb and new_proverb.strip():
        try:
            # Kiểm tra xem thành ngữ/tục ngữ đã tồn tại chưa
            if new_proverb in PROVERBS:
                st.warning(f"⚠️ Thành ngữ/tục ngữ '{new_proverb}' đã tồn tại trong từ điển!")
            else:
                # Thêm thành ngữ/tục ngữ mới vào từ điển bằng dictionary manager
                success, message = dict_manager.add_proverb(new_proverb, proverb_score)
                if success:
                    # Cập nhật biến toàn cục
                    reload_dictionaries()
                    
                    st.success(f"✅ Đã thêm thành ngữ/tục ngữ '{new_proverb}' với điểm số {proverb_score} (loại: {proverb_type}) vào từ điển!")
                    st.info("ℹ️ Thành ngữ/tục ngữ mới đã được thêm vào và sẵn sàng sử dụng ngay lập tức.")
                else:
                    st.error(f"❌ Không thể thêm thành ngữ/tục ngữ vào từ điển: {message}")
        except Exception as e:
            st.error(f"❌ Có lỗi xảy ra khi cập nhật từ điển: {str(e)}")
    search_proverb = st.text_input("Tìm kiếm thành ngữ/tục ngữ:", placeholder="Nhập thành ngữ/tục ngữ cần tìm...", key="search_proverb")
    try:
        reload_dictionaries()
        updated_proverbs = dict_manager.get_proverbs()
    except Exception:
        updated_proverbs = PROVERBS
    # Lọc và hiển thị từ điển
    filtered_proverbs = {k: v for k, v in updated_proverbs.items() if not search_proverb or search_proverb.lower() in k.lower()}
    
    st.subheader("Danh sách thành ngữ/tục ngữ hiện có")
    st.write(f"Tổng số thành ngữ/tục ngữ trong từ điển: **{len(filtered_proverbs)}**")
    
    if filtered_proverbs:
        # Tạo bảng hiển thị dữ liệu
        proverb_table = []
        for word, score in filtered_proverbs.items():
            sentiment = "Tích cực 🟢" if score >= 0.1 else "Tiêu cực 🔴" if score <= -0.1 else "Trung lập ⚪"
            proverb_table.append({"Thành ngữ/tục ngữ": word, "Điểm số": f"{score:.1f}", "Loại": sentiment})
        st.table(proverb_table)

# Tab thống kê lịch sử và xuất báo cáo
tab3 = st.tabs(["Thống kê & Báo cáo"])[0]
with tab3:
    st.header('Lịch sử phân tích & Xuất báo cáo')
    if st.session_state.analysis_history:
        df_history = pd.DataFrame([
            {'Thời gian': h['timestamp'], 'Văn bản': h['text'], 'Trạng thái': h.get('status', 'Không xác định')} for h in st.session_state.analysis_history
        ])
        st.dataframe(df_history)
        if st.button('Xuất báo cáo DOCX'):
            file_path = export_report_docx(st.session_state.analysis_history)
            with open(file_path, 'rb') as f:
                st.download_button('Tải báo cáo DOCX', f, file_name='bao_cao_phan_tich.docx')
    else:
        st.info('Chưa có lịch sử phân tích nào.')
